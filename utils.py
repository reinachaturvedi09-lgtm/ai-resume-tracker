"""
Core utilities module for resume parsing and ATS optimization.
Provides functions for PDF/DOCX parsing, text processing, and matching algorithm.
"""

import re
import string
from typing import Tuple, List, Dict, Optional
from pathlib import Path
import logging

import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


class ResumeParser:
    """
    Handles parsing of resume files in PDF and DOCX formats.
    """

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to the PDF file.
            
        Returns:
            Extracted text from the PDF.
            
        Raises:
            ValueError: If file cannot be read or is corrupted.
        """
        try:
            text = ""
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text += page.extract_text()
            return text
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise ValueError(f"Failed to parse PDF file: {str(e)}")

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to the DOCX file.
            
        Returns:
            Extracted text from the DOCX.
            
        Raises:
            ValueError: If file cannot be read or is corrupted.
        """
        try:
            text = ""
            document = Document(file_path)
            for paragraph in document.paragraphs:
                text += paragraph.text + "\n"
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
            return text
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise ValueError(f"Failed to parse DOCX file: {str(e)}")

    @staticmethod
    def parse_resume(file_path: str, file_type: str) -> str:
        """
        Parse resume file based on file type.
        
        Args:
            file_path: Path to the resume file.
            file_type: Type of file ('pdf' or 'docx').
            
        Returns:
            Extracted resume text.
            
        Raises:
            ValueError: If file type is not supported.
        """
        file_type = file_type.lower().strip('.')
        
        if file_type == 'pdf':
            return ResumeParser.parse_pdf(file_path)
        elif file_type == 'docx':
            return ResumeParser.parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}. Supported types: pdf, docx")


class TextPreprocessor:
    """
    Handles text preprocessing and cleaning.
    """

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean and normalize text.
        
        Args:
            text: Raw text to clean.
            
        Returns:
            Cleaned text.
        """
        # Convert to lowercase
        text = text.lower()
        
        # Remove URLs
        text = re.sub(r'http\S+|www\S+|https\S+', '', text, flags=re.MULTILINE)
        
        # Remove email addresses
        text = re.sub(r'\S+@\S+', '', text)
        
        # Remove special characters but keep spaces
        text = re.sub(r'[^\w\s]', ' ', text)
        
        # Remove extra whitespace
        text = ' '.join(text.split())
        
        return text

    @staticmethod
    def extract_keywords(text: str, min_length: int = 3) -> List[str]:
        """
        Extract meaningful keywords from text.
        
        Args:
            text: Text to extract keywords from.
            min_length: Minimum length of keywords.
            
        Returns:
            List of extracted keywords.
        """
        stopwords = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'be', 'been',
            'have', 'has', 'do', 'does', 'did', 'will', 'would', 'could', 'should',
            'may', 'might', 'can', 'this', 'that', 'these', 'those', 'i', 'you',
            'he', 'she', 'it', 'we', 'they', 'what', 'which', 'who', 'when',
            'where', 'why', 'how', 'all', 'each', 'every', 'both', 'few', 'more',
            'most', 'other', 'some', 'such', 'no', 'nor', 'not', 'only', 'same',
            'so', 'than', 'too', 'very', 'just', 'also', 'now', 'if', 'if', 'else'
        }
        
        words = text.split()
        keywords = [
            word for word in words
            if len(word) >= min_length and word not in stopwords
        ]
        
        return keywords


class ATSMatcher:
    """
    Implements matching algorithm to compare resume against job description.
    """

    def __init__(self):
        """
        Initialize the TF-IDF vectorizer.
        """
        self.vectorizer = TfidfVectorizer(
            max_features=500,
            stop_words='english',
            ngram_range=(1, 2),
            lowercase=True
        )
        self.job_description = ""
        self.resume_text = ""
        self.similarity_score = 0.0

    def calculate_match_score(
        self,
        resume_text: str,
        job_description: str
    ) -> float:
        """
        Calculate TF-IDF cosine similarity between resume and job description.
        
        Args:
            resume_text: Cleaned resume text.
            job_description: Cleaned job description text.
            
        Returns:
            Similarity score between 0 and 1.
        """
        self.resume_text = resume_text
        self.job_description = job_description
        
        try:
            vectors = self.vectorizer.fit_transform([job_description, resume_text])
            similarity = cosine_similarity(vectors[0:1], vectors[1:2])[0][0]
            self.similarity_score = float(similarity)
            
            return self.similarity_score
        except Exception as e:
            logger.error(f"Error calculating match score: {str(e)}")
            return 0.0

    def extract_matched_keywords(
        self,
        resume_text: str,
        job_description: str,
        top_n: int = 15
    ) -> Tuple[List[str], List[str]]:
        """
        Extract matched and missing keywords.
        
        Args:
            resume_text: Cleaned resume text.
            job_description: Cleaned job description text.
            top_n: Number of top keywords to return.
            
        Returns:
            Tuple of (matched_keywords, missing_keywords).
        """
        try:
            job_keywords = set(TextPreprocessor.extract_keywords(job_description))
            resume_keywords = set(TextPreprocessor.extract_keywords(resume_text))
            
            matched = list(job_keywords & resume_keywords)
            missing = list(job_keywords - resume_keywords)
            
            matched = sorted(matched, key=lambda x: len(x), reverse=True)[:top_n]
            missing = sorted(missing, key=lambda x: len(x), reverse=True)[:top_n]
            
            return matched, missing
        except Exception as e:
            logger.error(f"Error extracting keywords: {str(e)}")
            return [], []

    def get_match_percentage(self) -> float:
        """
        Get match percentage (0-100).
        
        Returns:
            Match percentage.
        """
        return round(self.similarity_score * 100, 2)


class ResumeOptimizer:
    """
    Provides recommendations for resume optimization.
    """

    @staticmethod
    def get_improvement_recommendations(
        match_score: float,
        missing_keywords: List[str],
        matched_keywords: List[str]
    ) -> Dict[str, any]:
        """
        Generate actionable recommendations based on match analysis.
        
        Args:
            match_score: Overall match score (0-100).
            missing_keywords: List of missing keywords.
            matched_keywords: List of matched keywords.
            
        Returns:
            Dictionary with recommendations and insights.
        """
        recommendations = {
            'overall_feedback': '',
            'priority_actions': [],
            'optimization_score': match_score,
            'keyword_coverage': len(matched_keywords) / (len(matched_keywords) + len(missing_keywords)) * 100
                if (len(matched_keywords) + len(missing_keywords)) > 0 else 0
        }

        if match_score >= 75:
            recommendations['overall_feedback'] = (
                "✅ Excellent Match! Your resume is well-aligned with the job description. "
                "Focus on highlighting your achievements and quantifiable results."
            )
        elif match_score >= 50:
            recommendations['overall_feedback'] = (
                "⚠️ Good Match! Your resume has relevant skills, but there are gaps. "
                "Consider adding the missing keywords naturally throughout your resume."
            )
        else:
            recommendations['overall_feedback'] = (
                "❌ Poor Match. Significant skill gaps detected. "
                "Review the job description carefully and consider tailoring your resume accordingly."
            )

        if missing_keywords:
            recommendations['priority_actions'].append(
                f"Add these key skills to your resume: {', '.join(missing_keywords[:5])}"
            )
        
        if match_score < 50:
            recommendations['priority_actions'].append(
                "Consider rewriting sections to include job description keywords naturally."
            )
        
        if len(matched_keywords) > 0:
            recommendations['priority_actions'].append(
                f"Highlight these matched skills prominently: {', '.join(matched_keywords[:3])}"
            )

        return recommendations


def full_analysis(
    resume_text: str,
    job_description: str,
    top_n_keywords: int = 15
) -> Dict[str, any]:
    """
    Perform complete analysis of resume against job description.
    
    Args:
        resume_text: Raw resume text.
        job_description: Raw job description text.
        top_n_keywords: Number of top keywords to extract.
        
    Returns:
        Comprehensive analysis dictionary.
    """
    cleaned_resume = TextPreprocessor.clean_text(resume_text)
    cleaned_job_desc = TextPreprocessor.clean_text(job_description)
    
    matcher = ATSMatcher()
    match_score = matcher.calculate_match_score(cleaned_resume, cleaned_job_desc)
    
    matched_keywords, missing_keywords = matcher.extract_matched_keywords(
        cleaned_resume,
        cleaned_job_desc,
        top_n=top_n_keywords
    )
    
    recommendations = ResumeOptimizer.get_improvement_recommendations(
        matcher.get_match_percentage(),
        missing_keywords,
        matched_keywords
    )
    
    return {
        'match_percentage': matcher.get_match_percentage(),
        'matched_keywords': matched_keywords,
        'missing_keywords': missing_keywords,
        'recommendations': recommendations,
        'resume_length': len(cleaned_resume.split()),
        'job_desc_length': len(cleaned_job_desc.split())
    }
